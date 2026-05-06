import re
from dataclasses import dataclass
from typing import List, Optional

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:  # pragma: no cover - local fallback for minimal environments
    @dataclass
    class _Document:
        page_content: str
        metadata: dict

    class RecursiveCharacterTextSplitter:
        def __init__(
            self,
            chunk_size: int = 512,
            chunk_overlap: int = 50,
            separators: Optional[List[str]] = None,
            length_function=len,
            add_start_index: bool = True,
        ):
            self.chunk_size = chunk_size
            self.chunk_overlap = chunk_overlap
            self.separators = separators or ["\n\n", "\n", " ", ".", ",", "\u200b"]
            self.length_function = length_function
            self.add_start_index = add_start_index

        def _split_text(self, text: str) -> List[tuple[str, int]]:
            if not text:
                return []

            chunks: List[tuple[str, int]] = []
            start = 0
            step = max(1, self.chunk_size - self.chunk_overlap)

            while start < len(text):
                end = min(len(text), start + self.chunk_size)
                window = text[start:end]

                split_at = -1
                for separator in self.separators:
                    idx = window.rfind(separator)
                    if idx > split_at:
                        split_at = idx

                if 0 < split_at < len(window) - 1:
                    end = start + split_at + 1
                    window = text[start:end]

                chunks.append((window, start))
                if end >= len(text):
                    break
                start = max(end - self.chunk_overlap, start + step)

            return chunks

        def create_documents(self, texts, metadatas=None):
            metadatas = metadatas or [{} for _ in texts]
            documents = []
            for text, metadata in zip(texts, metadatas):
                for chunk_text, start_index in self._split_text(text):
                    chunk_metadata = metadata.copy()
                    if self.add_start_index:
                        chunk_metadata["start_index"] = start_index
                    documents.append(_Document(page_content=chunk_text, metadata=chunk_metadata))
            return documents

@dataclass
class Chunk:
    text: str
    index: int
    start_char: int
    end_char: int
    metadata: dict

class DocumentChunker:
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        separators: Optional[List[str]] = None,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        if separators is None:
            separators = [
                "\n\n",
                "\n",
                " ",
                ".",
                ",",
                "\u200b",
            ]

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            length_function=len,
            add_start_index=True,
        )

    def chunk_text(self, text: str, metadata: Optional[dict] = None) -> List[Chunk]:
        docs = self.splitter.create_documents([text], metadatas=[metadata.copy() if metadata else {}])

        chunks = []

        for i, doc in enumerate(docs):
            start_char = int(doc.metadata.get("start_index", 0))
            end_char = start_char + len(doc.page_content)

            chunks.append(Chunk(
                text=doc.page_content,
                index=i,
                start_char=start_char,
                end_char=end_char,
                metadata=doc.metadata.copy()
            ))

        return chunks

    def chunk_file(self, file_path: str, metadata: Optional[dict] = None) -> List[Chunk]:
        from pathlib import Path

        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".txt":
            text = path.read_text(encoding="utf-8")
        elif ext == ".md":
            text = path.read_text(encoding="utf-8")
        elif ext == ".pdf":
            text = self._extract_pdf_text(file_path)
        elif ext == ".docx":
            text = self._extract_docx_text(file_path)
        else:
            text = path.read_text(encoding="utf-8")

        file_meta = metadata.copy() if metadata else {}
        file_meta.update({
            "source": str(path),
            "filename": path.name,
        })

        return self.chunk_text(text, file_meta)

    def _extract_pdf_text(self, file_path: str) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            print(f"PDF extraction failed: {e}")
            return ""

    def _extract_docx_text(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return ""

def benchmark_chunking_strategies(
    text: str,
    strategies: List[dict],
) -> dict:
    results = {}

    for strategy in strategies:
        chunker = DocumentChunker(
            chunk_size=strategy.get("chunk_size", 512),
            chunk_overlap=strategy.get("chunk_overlap", 50),
            separators=strategy.get("separators"),
        )

        chunks = chunker.chunk_text(text)

        results[strategy["name"]] = {
            "chunk_size": strategy.get("chunk_size", 512),
            "chunk_overlap": strategy.get("chunk_overlap", 50),
            "separators": strategy.get("separators"),
            "num_chunks": len(chunks),
            "avg_chunk_size": sum(len(c.text) for c in chunks) / len(chunks) if chunks else 0,
            "chunks": chunks,
        }

    return results
